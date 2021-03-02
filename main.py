import os
from os.path import isfile, join
from matplotlib import pyplot as plt
from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
import pandas as pd
import numpy as np
import matplotlib.dates as mdates
from matplotlib.patches import Rectangle
import math
import warnings

cm = 1/2.54  # centimeters in inches
exclude_rows = [1, 2, 3, 4, 5, 6]  # rows to exclude in reading CSV file
graph_types = ["Current [A]", "THD-F [%]", "Power factor", "Voltage [V]"]  # y axis
# periods = ["weekdays", "all", "saturdays"] #measurement period/x axis in graph


def sitename():
    sitename = input("Enter the site shortform: ").upper()

    return sitename


def save_to_png(plt, file, opt):
	script_dir = os.path.dirname(__file__)
	results_dir = os.path.join(script_dir, 'graphs/')
	if not os.path.isdir(results_dir):
		os.makedirs(results_dir)
	fig = plt.savefig('graphs/'+opt+"_"+os.path.splitext(file)[0]+'.png', bbox_inches='tight', pad_inches=1)
	plt.close(fig)


def graph_meas_period(columns, opt, file):
	# plot graph
	new_df = df[["FormattedTime"] + columns]
	idx = new_df.index
	num_of_rows = len(idx)
	new_df.plot(x="FormattedTime", y=columns, figsize=(25 * 3 * cm, 5 * 3 * cm), color=['black', 'blue', 'red'],
				x_compat=True)

	if opt != "Power factor":
		legend = plt.legend(bbox_to_anchor=(0.65, -0.1), ncol=3, columnspacing=15)
		# set legend text
		for col in columns:
			if "1" in col:
				color = 'black'
			elif "2" in col:
				color = 'blue'
			elif "3" in col:
				color = 'red'
			else:
				color = 'black'

			legend.get_texts()[columns.index(col)].set_text(opt.split(" ")[0] + " " + col)
			legend.legendHandles[columns.index(col)].set_color(color)
	if opt == "Power factor":
		plt.ylim([0, 1.1])

	if opt == "Voltage [V]":
		plt.ylim(0, 250)

	# setup graph
	plt.xlim(df["FormattedTime"].iloc[0], df["FormattedTime"].iloc[-1])

	hours = mdates.HourLocator(interval=6)
	h_fmt = mdates.DateFormatter('%Y-%m-%d\n%H:%M:%S')
	plt.gca().xaxis.set_major_locator(hours)
	plt.gca().xaxis.set_major_formatter(h_fmt)

	plt.xticks(fontsize=6)
	plt.grid()
	plt.xlabel("")
	plt.ylabel(opt)

	save_to_png(plt, file, opt)
	return num_of_rows


def graph_additional_current(opt, columns, file):
	new_df = df[["Date"] + ["FormattedTime"] + ["dateTime"] + columns]
	if opt == "Weekday":
		new_df = new_df[new_df["FormattedTime"].dt.dayofweek < 5]
	elif opt == "Weekend":
		new_df = new_df[new_df["FormattedTime"].dt.dayofweek >= 5]

	new_df.plot(x="Date", y=columns, figsize=(25*3*cm, 5*3*cm), color=['black', 'blue', 'red'],x_compat=True)
	legend = plt.legend(bbox_to_anchor=(0.675, -0.15), ncol=3, columnspacing=15)

	# set legend text
	for col in columns:
		if "1" in col:
			color = 'black'
		elif "2" in col:
			color = 'blue'
		elif "3" in col:
			color = 'red'
		else:
			color = 'black'

		legend.get_texts()[columns.index(col)].set_text("Current " + col)
		legend.legendHandles[columns.index(col)].set_color(color)

	plt.xticks(fontsize=6)
	plt.grid()
	plt.xlabel("")
	plt.ylabel("Current [A]")

	save_to_png(plt, file, "Current_"+opt)


def graph_cur_vs_thd(cols_current, cols_thd):
	graph_df = df[cols_current + cols_thd]
	ax1 = graph_df.plot(kind='scatter', x="I1[A]", y="THD-F_I1[%]", figsize=(15 * 3 * cm, 5 * 3 * cm),
						label="Phase 1", color='black')
	ax2 = graph_df.plot(kind='scatter', x="I2[A]", y="THD-F_I2[%]", figsize=(15 * 3 * cm, 5 * 3 * cm),
						label="Phase 2", ax=ax1, color='b')
	ax3 = graph_df.plot(kind='scatter', x="I3[A]", y="THD-F_I3[%]", figsize=(15 * 3 * cm, 5 * 3 * cm),
						label="Phase 3", ax=ax1, color='r')

	# setup graph
	min_current = graph_df[["I1[A]", "I2[A]", "I3[A]"]].min().min()
	min_multiplier = min_current if (graph_df[["I1[A]", "I2[A]", "I3[A]"]].min().min() < 0) else 0
	max_current = graph_df[["I1[A]", "I2[A]", "I3[A]"]].max().max()
	max_multiplier = max_current if (graph_df[["I1[A]", "I2[A]", "I3[A]"]].max().max() > 2000) else 0

	# plt.xlim(min_current,max_current)
	plt.ylim(0, graph_df[["THD-F_I1[%]", "THD-F_I2[%]", "THD-F_I3[%]"]].max().max())
	plt.xticks(fontsize=6)
	plt.grid()
	plt.xlabel("Phase Current (A)")
	plt.ylabel("Current Harmonics (%)")

	# shade scatter plot
	plt.gca().add_patch(Rectangle((-100 + (min_multiplier), 0), 140 + (min_multiplier * -1), 20,
								  color='grey', alpha=0.5))
	plt.gca().add_patch(Rectangle((40, 0), 360, 15,
								  color='grey', alpha=0.5))
	plt.gca().add_patch(Rectangle((400, 0), 400, 12,
								  color='grey', alpha=0.5))
	plt.gca().add_patch(Rectangle((800, 0), 1200, 8,
								  color='grey', alpha=0.5))
	plt.gca().add_patch(Rectangle((2000, 0), 1000 + max_multiplier, 5,
								  color='grey', alpha=0.5))

	save_to_png(plt, file, "Current_vs_THD")


def get_columns(opt):
	switcher = {
		"Current [A]": 	["I1[A]", "I2[A]", "I3[A]"],
		"THD-F [%]": 	["THD-F_I1[%]", "THD-F_I2[%]", "THD-F_I3[%]"],
		"Power factor": ["PF"],
		"Voltage [V]": 	["U1[V]", "U2[V]", "U3[V]"]
	}
	return switcher.get(opt, "Invalid input")


def get_df(file):  # get dataFrame from csv
	data = pd.read_csv(file, skiprows=lambda x: x in exclude_rows, header=0)
	new_data = data.assign(dateTime=data.Date + "\n" + data.Time, FormattedTime=data.Date + " " + data.Time)
	new_data['FormattedTime'] = pd.to_datetime(new_data['FormattedTime'], format="%d/%m/%Y %H:%M:%S")
	return new_data


def get_max(df, col):
	max_index = df[col].idxmax()
	return df[col].max(), df.iloc[max_index, 0] + ' ' + df.iloc[max_index, 1]


def get_min(df, col):
	min_index = df[col].idxmin()
	return df[col].min(), df.iloc[min_index, 0] + ' ' + df.iloc[min_index, 1]


def get_circuit_operation(df):
	sLength = len(df['Date'])
	max_i1 = round(df['I1[A]'].max(), 2)
	max_i2 = round(df['I2[A]'].max(), 2)
	max_i3 = round(df['I3[A]'].max(), 2)
	max_i = min(max_i2, max_i2, max_i3)
	min_i1 = round(df['I1[A]'].min(), 2)
	min_i2 = round(df['I2[A]'].min(), 2)
	min_i3 = round(df['I3[A]'].min(), 2)
	min_i = max(min_i1, min_i2, min_i3)
	average_i = (max_i - min_i)*0.5
	df['Weekday'] = pd.Series(np.random.randn(sLength), index=df.index)
	df['Hour'] = pd.Series(np.random.randn(sLength), index=df.index)
	df['dateTime'] = pd.Series(np.random.randn(sLength), index=df.index)
	df['Minute'] = pd.Series(np.random.randn(sLength), index=df.index)
	df['dateTime'] = df['Date'] + ' ' + df['Time']
	df['dateTime'] = pd.to_datetime(df['dateTime'], format="%d/%m/%Y %H:%M:%S")
	df['Hour'] = df['dateTime'].dt.hour
	df['Minute'] = df['dateTime'].dt.minute
	df['Weekday'] = df['dateTime'].dt.dayofweek
	base_load_mon = []
	base_load_tue = []
	base_load_wed = []
	base_load_thu = []
	base_load_fri = []
	base_load_sat = []
	base_load_sun = []
	weekday_max_base = 0
	for i in range(len(df["Hour"])-2):
		if df['Weekday'].iloc[i] == 0:
			if df["I1[A]"].iloc[i] > average_i or df["I2[A]"].iloc[i] > average_i or df["I3[A]"].iloc[i] > average_i:
				base_load_mon.append(df['Time'].iloc[i])
			elif df["I1[A]"].iloc[i] < average_i and df["I1[A]"].iloc[i-2] < average_i and df["I1[A]"].iloc[i+2] < average_i:
				if df["I1[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I1[A]"].iloc[i]
			elif df["I2[A]"].iloc[i] < average_i and df["I2[A]"].iloc[i-2] < average_i and df["I2[A]"].iloc[i+2] < average_i:
				if df["I2[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I2[A]"].iloc[i]
			elif df["I3[A]"].iloc[i] < average_i and df["I3[A]"].iloc[i-2] < average_i and df["I3[A]"].iloc[i+2] < average_i:
				if df["I3[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I3[A]"].iloc[i]
		elif df['Weekday'].iloc[i] == 1:
			if df["I1[A]"].iloc[i] > average_i or df["I2[A]"].iloc[i] > average_i or df["I3[A]"].iloc[i] > average_i:
				base_load_tue.append(df['Time'].iloc[i])
			elif df["I1[A]"].iloc[i] < average_i and df["I1[A]"].iloc[i - 2] < average_i and df["I1[A]"].iloc[i + 2] < average_i:
				if df["I1[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I1[A]"].iloc[i]
			elif df["I2[A]"].iloc[i] < average_i and df["I2[A]"].iloc[i - 2] < average_i and df["I2[A]"].iloc[i + 2] < average_i:
				if df["I2[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I2[A]"].iloc[i]
			elif df["I3[A]"].iloc[i] < average_i and df["I3[A]"].iloc[i - 2] < average_i and df["I3[A]"].iloc[i + 2] < average_i:
				if df["I3[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I3[A]"].iloc[i]
		elif df['Weekday'].iloc[i] == 2:
			if df["I1[A]"].iloc[i] > average_i or df["I2[A]"].iloc[i] > average_i or df["I3[A]"].iloc[i] > average_i:
				base_load_wed.append(df['Time'].iloc[i])
			elif df["I1[A]"].iloc[i] < average_i and df["I1[A]"].iloc[i - 2] < average_i and df["I1[A]"].iloc[i + 2] < average_i:
				if df["I1[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I1[A]"].iloc[i]
			elif df["I2[A]"].iloc[i] < average_i and df["I2[A]"].iloc[i - 2] < average_i and df["I2[A]"].iloc[i + 2] < average_i:
				if df["I2[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I2[A]"].iloc[i]
			elif df["I3[A]"].iloc[i] < average_i and df["I3[A]"].iloc[i - 2] < average_i and df["I3[A]"].iloc[i + 2] < average_i:
				if df["I3[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I3[A]"].iloc[i]
		elif df['Weekday'].iloc[i] == 3:
			if df["I1[A]"].iloc[i] > average_i or df["I2[A]"].iloc[i] > average_i or df["I3[A]"].iloc[i] > average_i:
				base_load_thu.append(df['Time'].iloc[i])
			elif df["I1[A]"].iloc[i] < average_i and df["I1[A]"].iloc[i - 2] < average_i and df["I1[A]"].iloc[i + 2] < average_i:
				if df["I1[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I1[A]"].iloc[i]
			elif df["I2[A]"].iloc[i] < average_i and df["I2[A]"].iloc[i - 2] < average_i and df["I2[A]"].iloc[i + 2] < average_i:
				if df["I2[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I2[A]"].iloc[i]
			elif df["I3[A]"].iloc[i] < average_i and df["I3[A]"].iloc[i - 2] < average_i and df["I3[A]"].iloc[i + 2] < average_i:
				if df["I3[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I3[A]"].iloc[i]
		elif df['Weekday'].iloc[i] == 4:
			if df["I1[A]"].iloc[i] > average_i or df["I2[A]"].iloc[i] > average_i or df["I3[A]"].iloc[i] > average_i:
				base_load_fri.append(df['Time'].iloc[i])
			elif df["I1[A]"].iloc[i] < average_i and df["I1[A]"].iloc[i - 2] < average_i and df["I1[A]"].iloc[i + 2] < average_i:
				if df["I1[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I1[A]"].iloc[i]
			elif df["I2[A]"].iloc[i] < average_i and df["I2[A]"].iloc[i - 2] < average_i and df["I2[A]"].iloc[i + 2] < average_i:
				if df["I2[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I2[A]"].iloc[i]
			elif df["I3[A]"].iloc[i] < average_i and df["I3[A]"].iloc[i - 2] < average_i and df["I3[A]"].iloc[i + 2] < average_i:
				if df["I3[A]"].iloc[i] > weekday_max_base:
					weekday_max_base = df["I3[A]"].iloc[i]
		elif df['Weekday'].iloc[i] == 5:
			if df["I1[A]"].iloc[i] > average_i or df["I2[A]"].iloc[i] > average_i or df["I3[A]"].iloc[i] > average_i:
				base_load_sat.append(df['Time'].iloc[i])
		elif df['Weekday'].iloc[i] == 6:
			if df["I1[A]"].iloc[i] > average_i or df["I2[A]"].iloc[i] > average_i or df["I3[A]"].iloc[i] > average_i:
				base_load_sun.append(df['Time'].iloc[i])
	base_load_start_weekday = min(base_load_mon[0], base_load_tue[0], base_load_wed[0], base_load_thu[0],
								  base_load_fri[0])
	base_load_end_weekday = max(base_load_mon[-1], base_load_tue[-1], base_load_wed[-1], base_load_thu[-1],
								base_load_fri[-1])
	weekday_max_base = (math.ceil(weekday_max_base/10))*10
	if len(base_load_sun) < 4:
		detail = {
			"Weekday start": base_load_start_weekday,
			"Weekday end": base_load_end_weekday,
			"Sat start": base_load_sat[0],
			"Sat end": base_load_sat[-1],
			"Sun start": 0,
			"Sun end": 0,
			"Weekday baseload": weekday_max_base,
		}
	if len(base_load_sun) > 3:
		detail = {
			"Weekday start": base_load_start_weekday,
			"Weekday end": base_load_end_weekday,
			"Sat start": base_load_sat[0],
			"Sat end": base_load_sat[-1],
			"Sun start": base_load_sun[0],
			"Sun end": base_load_sun[-1],
			"Weekday baseload": weekday_max_base,
		}
	return detail


def get_average(df):
	maxi_THD1 = round(df['THD-F_I1[%]'].max(), 2)
	maxi_THD2 = round(df['THD-F_I2[%]'].max(), 2)
	maxi_THD3 = round(df['THD-F_I3[%]'].max(), 2)
	average_THD1 = round(df['THD-F_I1[%]'].mean(skipna=True), 2)
	average_THD2 = round(df['THD-F_I2[%]'].mean(skipna=True), 2)
	average_THD3 = round(df['THD-F_I3[%]'].mean(skipna=True), 2)
	ave = (average_THD1 + average_THD2 + average_THD3) / 3
	average_THD = round(ave, 2)
	average_i1 = round(df['I1[A]'].mean(skipna=True), 2)
	average_i2 = round(df['I2[A]'].mean(skipna=True), 2)
	average_i3 = round(df['I3[A]'].mean(skipna=True), 2)
	exchange = min(df['I1[A]'].min(), df['I2[A]'].min(), df['I3[A]'].min())
	min_i = round(exchange, 2)
	average_PF = round(df['PF'].mean(skipna=True), 2)
	max_pf = round(df['PF'].max(), 2)
	min_pf = round(df['PF'].min(), 2)
	average_v1 = round(df['U1[V]'].mean(skipna=True), 2)
	average_v2 = round(df['U2[V]'].mean(skipna=True), 2)
	average_v3 = round(df['U3[V]'].mean(skipna=True), 2)
	exchange = (maxi_THD1+maxi_THD2+maxi_THD3)/3
	max_thd = round(exchange, 2)
	details = {
		"AVE i1": average_i1,
		"AVE i2": average_i2,
		"AVE i3": average_i3,
		"MIN i": min_i,
		"AVE v1": average_v1,
		"AVE v2": average_v2,
		"AVE v3": average_v3,
		"AVE PF": average_PF,
		"PF MAX": max_pf,
		"PF MIN": min_pf,
		"AVE THD1": average_THD1,
		"AVE THD2": average_THD2,
		"AVE THD3": average_THD3,
		"MAX THD1": maxi_THD1,
		"MAX THD2": maxi_THD2,
		"MAX THD3": maxi_THD3,
		"MAX THD": max_thd,
		"AVE THD": average_THD,
	}
	return details

def get_circuit_details(data, file):
	details = {
			"Circuit Name": os.path.splitext(file)[0].split("_")[1],
			"Circuit current": os.path.splitext(file)[0].split("_")[2],
			"THD 1 MAX": get_max(data, "THD-F_I1[%]"),
			"THD 1 MIN": get_min(data, "THD-F_I1[%]"),
			"THD 2 MAX": get_max(data, "THD-F_I2[%]"),
			"THD 2 MIN": get_min(data, "THD-F_I2[%]"),
			"THD 3 MAX": get_max(data, "THD-F_I3[%]"),
			"THD 3 MIN": get_min(data, "THD-F_I3[%]"),

			"CUR 1 MAX": get_max(data, "I1[A]"),
			"CUR 1 MIN": get_min(data, "I1[A]"),
			"CUR 2 MAX": get_max(data, "I2[A]"),
			"CUR 2 MIN": get_min(data, "I2[A]"),
			"CUR 3 MAX": get_max(data, "I3[A]"),
			"CUR 3 MIN": get_min(data, "I3[A]"),

			"VOL 1 MAX": get_max(data, "U1[V]"),
			"VOL 1 MIN": get_min(data, "U1[V]"),
			"VOL 2 MAX": get_max(data, "U2[V]"),
			"VOL 2 MIN": get_min(data, "U2[V]"),
			"VOL 3 MAX": get_max(data, "U3[V]"),
			"VOL 3 MIN": get_min(data, "U3[V]"),

			"PF MAX": get_max(data, "PF"),
			"PF MIN": get_min(data, "PF"),
	}
	return details


def add_graph_to_document(document, circuit, opt, idx):
	type_map= {
				"Current [A]": "Current",
				"THD-F [%]": "Current of THD-F",
				"Power factor": "Power factor",
				"Voltage [V]": "Voltage"
	}

	filepath = os.path.abspath('graphs/' + opt + '_' + str(idx+1) + '_' + circuit['Circuit Name'] + '_' + circuit["Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph(type_map.get(opt) + ' of ' + circuit["Circuit Name"] + ' for Measurement Period ')
	document.add_paragraph()

	if opt == "Current [A]":
		filepath = os.path.abspath('graphs/Current_Weekday_' + str(idx+1) + '_' + circuit['Circuit Name'] + '_' + circuit["Circuit current"] + '.png')
		document.add_picture(filepath, width=Inches(6.0))
		document.add_paragraph('Current of ' + circuit["Circuit Name"] + ' for Typical Weekday')
		document.add_paragraph()

		filepath = os.path.abspath('graphs/Current_Weekend_' + str(idx+1) + '_' + circuit['Circuit Name'] + '_' + circuit["Circuit current"] + '.png')
		document.add_picture(filepath, width=Inches(6.0))
		document.add_paragraph('Current of ' + circuit["Circuit Name"] + ' for Typical Saturday and Sunday')
		document.add_paragraph()


def power_quality_appendix(circuits, circuit_total):
	document = Document()
	document.add_heading('Appendix D — LV Switchboard Power Quality', 0)
	for c in circuits:
		for g in graph_types:
			add_graph_to_document(document, c, g, circuits.index(c))
	document.save('Appendix D.docx')


number = ["zero", "one", "two", "three", "four", "five", "six", "seven", "eight", "nine", "ten",
		   "eleven", "twelve", "thirteen", "fourteen", "fifteen", "sixteen", "seventeen", "eighteen", "nineteen"]


# MAIN #
circuit_total = 0  # total number of circuit (excel)
incomer_total = 0  # total number of incomer (incomer will container in the excel name)
current_details = {}  # each current name and their rated current (e.g. kitty incomer_3200)
number_rows = []
site_name = sitename()

circuit_details = list()
circuit_average = list()
incomer_detail = list()
# max and min and their coorsponding time or current, voltage, power factor and THD
# average of current, voltage, power factor and THD

warnings.simplefilter("ignore")
path = 'files/'
# LOOP for each circuit/excel file
for file in os.listdir(path):
	if isfile(join(path, file)):
		circuit_total += 1
		count = 0
		test = "Incomer"
		count = file.count(test)
		df = get_df(join(path, file))
		df['PF'] = df['PF'].abs()
		circuit_details.append(get_circuit_details(df, file))
		circuit_average.append(get_average(df))

		if count == 1:
			incomer_total += 1
			incomer_detail.append(get_circuit_operation(df))

		for option in graph_types:  # loop to graph each type
			col_list = get_columns(option)
			num_rows = graph_meas_period(col_list, option, file)
			number_rows.append(str(num_rows))
		graph_cur_vs_thd(get_columns("Current [A]"), get_columns("THD-F [%]"))
		graph_additional_current("Weekday", get_columns("Current [A]"), file)
		graph_additional_current("Weekend", get_columns("Current [A]"), file)

# write circuit info to file
# circuits_df = pd.DataFrame.from_dict(circuit_details)
# circuits_df.to_excel('circuit_details.xlsx')

print(f'Number of circuits: {circuit_total}\nTotal Incomer: {incomer_total}\n')
range(len(circuit_details))
range(len(circuit_average))
numbers = str(incomer_total)
circuit400 = circuit_total - incomer_total
incomer_tot = str(incomer_total)
circuit_number = str(circuit400)

document = Document()
document.add_heading('Power Quality Analysis', 0)
p = document.add_paragraph(site_name)
p.add_run(' is served by ')
p.add_run(number[incomer_total] + ' ')
p.add_run(circuit_details[0]["Circuit current"])
p.add_run('A incomers. ')
p.add_run('Energenz conducted power quality measurements on ')
p.add_run(number[incomer_total])
p.add_run(' (' + incomer_tot + ') ')
p.add_run('of the incomers and ')
p.add_run(number[circuit400])
p.add_run(' (' + circuit_number + ') ')
p.add_run('400A subcircuits:')

for x in range(circuit_total):
	p = document.add_paragraph(circuit_details[x]["Circuit current"], style='List Bullet')
	p.add_run('A - ')
	p.add_run(circuit_details[x]["Circuit Name"])

document.add_paragraph('The section below presented the voltage (V), current (A), power factor (PF), ' +
					   'and total harmonic distortion (THD) findings of each circuit.')

for x in range(incomer_total):
	document.add_heading(circuit_details[x]["Circuit Name"] + ' (' + circuit_details[x]["Circuit current"] + 'A)', level=1)
	p = document.add_paragraph().add_run('Current and Voltage')
	p.bold = True

	baseload_weekday_start = str(incomer_detail[x]["Weekday start"])
	if len(baseload_weekday_start) == 7:
		baseload_weekday_start = "0" + baseload_weekday_start
	baseload_weekday_start = baseload_weekday_start[:5]
	baseload_weekday_end = str(incomer_detail[x]["Weekday end"])
	if len(baseload_weekday_end) == 7:
		baseload_weekday_end = "0" + baseload_weekday_end
	baseload_weekday_end = baseload_weekday_end[:5]

	incomer_curmin = round(circuit_average[x]["MIN i"], -1)
	incomer_curmin = str(incomer_curmin)
	incomer_curmax = round(incomer_detail[x]["Weekday baseload"], -1)
	incomer_curmax = str(incomer_curmax)
	document.add_paragraph('The daily pattern of current demand is different from weekday to weekend. ' +
						   'The current demand in weekday increases at around ' + baseload_weekday_start +
						   ' and decreases at around ' + baseload_weekday_end + '. The baseload in weekday is ' +
						   'about '+ incomer_curmin + 'A to ' + incomer_curmax + 'A at non-peak hours')
	baseload_sat_start = str(incomer_detail[x]["Sat start"])
	if len(baseload_sat_start) == 7:
		baseload_sat_start = "0" + baseload_sat_start
	baseload_sat_start = baseload_sat_start[:5]
	baseload_sat_end = str(incomer_detail[x]["Sat end"])
	if len(baseload_sat_end) == 7:
		baseload_sat_end = "0" + baseload_sat_end
	baseload_sat_end = baseload_sat_end[:5]

	document.add_paragraph('The current demand on Saturday increases at around ' + baseload_sat_start +
						   ' and decreases at around ' + baseload_sat_end + '. The baseload on Saturday is ' +
						   'about '+ incomer_curmin + 'A to ' + incomer_curmax + 'A at non-peak hours')

	if incomer_detail[x]["Sun start"] == 0:
		document.add_paragraph('The current demand on Sunday stays within the base load of ' +  incomer_curmin +
							   'A to ' + incomer_curmax + 'A for the whole day.')
	else:
		baseload_sun_start = str(incomer_detail[x]["Sun start"])
		if len(baseload_sun_start) == 7:
			baseload_sun_start = "0" + baseload_sun_start
		baseload_sun_start = baseload_sun_start[:5]
		baseload_sun_end = str(incomer_detail[x]["Sun end"])
		if len(baseload_sun_end) == 7:
			baseload_sun_end = "0" + baseload_sun_end
		baseload_sun_end = baseload_sun_end[:5]
		document.add_paragraph('The current demand on Sunday increases at around ' + baseload_sun_start +
							   ' and decreases at around ' + baseload_sun_end + '. The baseload on Sunday is ' +
							   'about ' + incomer_curmin + 'A to ' + incomer_curmax + 'A at non-peak hours')

	document.add_paragraph('The current and voltage measurement are presented below.')

	# add current table and label format
	document.add_paragraph('Current of ' + circuit_details[x]["Circuit Name"] + ' for Measurement Period Data Summary')
	table = document.add_table(rows=4, cols=5)
	table.style = 'Light List Accent 5'
	label_cell0 = table.cell(0, 0)
	label0_paragraph = label_cell0.paragraphs[0]
	run0 = label0_paragraph.add_run('Parameter')
	run0.font.color.rgb = RGBColor(255, 255, 255)
	label_cell1 = table.cell(0, 1)
	label1_paragraph = label_cell1.paragraphs[0]
	run1 = label1_paragraph.add_run('Times')
	run1.font.color.rgb = RGBColor(255, 255, 255)
	label_cell2 = table.cell(0, 2)
	label2_paragraph = label_cell2.paragraphs[0]
	run2 = label2_paragraph.add_run('Current RMS\n Max.')
	run2.font.color.rgb = RGBColor(255, 255, 255)
	label_cell3 = table.cell(0, 3)
	label3_paragraph = label_cell3.paragraphs[0]
	run3 = label3_paragraph.add_run('Current RMS\n Min.')
	run3.font.color.rgb = RGBColor(255, 255, 255)
	label_cell4 = table.cell(0, 4)
	label4_paragraph = label_cell4.paragraphs[0]
	run4 = label4_paragraph.add_run('Current RMS Average')
	run4.font.color.rgb = RGBColor(255, 255, 255)

	# information of the circuit
	first_cells = table.rows[1].cells
	first_cells[0].text = circuit_details[x]["Circuit Name"] + ' I1 [A]'
	first_cells[1].text = number_rows[x]
	curmax = str(circuit_details[x]["CUR 1 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	first_cells[2].text = curmax
	curmin = str(circuit_details[x]["CUR 1 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	first_cells[3].text = curmin
	curave = str(circuit_average[x]["AVE i1"])
	first_cells[4].text = curave
	second_cells = table.rows[2].cells
	second_cells[0].text = circuit_details[x]["Circuit Name"] + ' I2 [A]'
	second_cells[1].text = number_rows[x]
	curmax = str(circuit_details[x]["CUR 2 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	second_cells[2].text = curmax
	curmin = str(circuit_details[x]["CUR 2 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	second_cells[3].text = curmin
	curave = str(circuit_average[x]["AVE i2"])
	second_cells[4].text = curave
	third_cells = table.rows[3].cells
	third_cells[0].text = circuit_details[x]["Circuit Name"] + ' I2 [A]'
	third_cells[1].text = number_rows[x]
	curmax = str(circuit_details[x]["CUR 3 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	third_cells[2].text = curmax
	curmin = str(circuit_details[x]["CUR 3 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	third_cells[3].text = curmin
	curave = str(circuit_average[x]["AVE i3"])
	third_cells[4].text = curave
	document.add_paragraph()

	# add current of first circuit
	i = x+1
	filepath = os.path.abspath('graphs/Current [A]_' + str(i) + '_' + circuit_details[x]["Circuit Name"] + '_' + circuit_details[x]["Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Current of ' + circuit_details[x]["Circuit Name"] + ' for Measurement Period ')
	document.add_paragraph()
	filepath = os.path.abspath(
		'graphs/Current_Weekday_' + str(i) + '_' + circuit_details[x]["Circuit Name"] + '_' + circuit_details[x][
			"Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Current of ' + circuit_details[x]["Circuit Name"] + ' for Typical Weekday ')
	document.add_paragraph()
	filepath = os.path.abspath(
		'graphs/Current_Weekend_' + str(i) + '_' + circuit_details[x]["Circuit Name"] + '_' + circuit_details[x][
			"Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Current of ' + circuit_details[x]["Circuit Name"] + ' for Typical Weekend ')
	document.add_paragraph()
	document.add_paragraph()

	# add voltage table and label format
	document.add_paragraph('Voltage of ' + circuit_details[x]["Circuit Name"] + ' for Measurement Period Data Summary')
	table1 = document.add_table(rows=4, cols=5)
	table1.style = 'Light List Accent 5'
	label_cell0 = table1.cell(0, 0)
	label0_paragraph = label_cell0.paragraphs[0]
	run0 = label0_paragraph.add_run('Parameter')
	run0.font.color.rgb = RGBColor(255, 255, 255)
	label_cell1 = table1.cell(0, 1)
	label1_paragraph = label_cell1.paragraphs[0]
	run1 = label1_paragraph.add_run('Times')
	run1.font.color.rgb = RGBColor(255, 255, 255)
	label_cell2 = table1.cell(0, 2)
	label2_paragraph = label_cell2.paragraphs[0]
	run2 = label2_paragraph.add_run('Voltage RMS\n Max.')
	run2.font.color.rgb = RGBColor(255, 255, 255)
	label_cell3 = table1.cell(0, 3)
	label3_paragraph = label_cell3.paragraphs[0]
	run3 = label3_paragraph.add_run('Voltage RMS\n Min.')
	run3.font.color.rgb = RGBColor(255, 255, 255)
	label_cell4 = table1.cell(0, 4)
	label4_paragraph = label_cell4.paragraphs[0]
	run4 = label4_paragraph.add_run('Voltage RMS Average')
	run4.font.color.rgb = RGBColor(255, 255, 255)

	# information of the circuit
	first_cells = table1.rows[1].cells
	first_cells[0].text = circuit_details[x]["Circuit Name"] + ' U1 [V]'
	first_cells[1].text = number_rows[x]
	curmax = str(circuit_details[x]["VOL 1 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	first_cells[2].text = curmax
	curmin = str(circuit_details[x]["VOL 1 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	first_cells[3].text = curmin
	curave = str(circuit_average[x]["AVE v2"])
	first_cells[4].text = curave
	second_cells = table1.rows[2].cells
	second_cells[0].text = circuit_details[x]["Circuit Name"] + ' U2 [V]'
	second_cells[1].text = number_rows[x]
	curmax = str(circuit_details[x]["VOL 2 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	second_cells[2].text = curmax
	curmin = str(circuit_details[x]["VOL 2 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	second_cells[3].text = curmin
	curave = str(circuit_average[x]["AVE v2"])
	second_cells[4].text = curave
	third_cells = table1.rows[3].cells
	third_cells[0].text = circuit_details[x]["Circuit Name"] + ' U2 [V]'
	third_cells[1].text = number_rows[x]
	curmax = str(circuit_details[x]["VOL 3 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	third_cells[2].text = curmax
	curmin = str(circuit_details[x]["VOL 3 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	third_cells[3].text = curmin
	curave = str(circuit_average[x]["AVE v3"])
	third_cells[4].text = curave

	# add current of first voltage
	filepath = os.path.abspath(
		'graphs/Voltage [V]_' + str(i) + '_' + circuit_details[x]["Circuit Name"] + '_' + circuit_details[x][
			"Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Voltage of ' + circuit_details[x]["Circuit Name"] + ' for Measurement Period')

	document.add_paragraph()
	p = document.add_paragraph().add_run('Total Harmonic Distortion')
	p.bold = True
	document.add_paragraph('The Code of Practice (COP) of Energy Efficiency of Building Services Installation has ' +
						   'regulated the maximum THD at difference rated circuit current. The details are shown below.')
	filepath = os.path.abspath('graphs/THD'+ '.png')
	document.add_picture(filepath, width=Inches(4.0))
	thd = document.add_paragraph().add_run("Please add the red box indicated the current range after delete this sentence.")
	thd.font.color.rgb = RGBColor(255, 0, 0)
	document.add_paragraph('Code of Practice of Energy Efficiency of Building Services Installation')

	if int(circuit_details[x]["Circuit current"])>= 400 and int(circuit_details[x]["Circuit current"]) < 800:
		circuitTHD = "12.0%"
		wording = "is between 400A and 800A"
		if int(circuit_average[x]["MAX THD"]) > 12:
			word = "does not fulfill"
		else:
			word = "fulfills"
	elif int(circuit_details[x]["Circuit current"]) >= 800 and int(circuit_details[x]["Circuit current"]) < 2000:
		circuitTHD = "8.0%"
		wording = "is between 800A and 2,000A"
		if int(circuit_average[x]["MAX THD"]) > 8:
			word = "does not fulfill"
		else:
			word = "fulfills"
	elif int(circuit_details[x]["Circuit current"]) >= 2000:
		circuitTHD = "5.0%"
		wording = "is greater than 2,000A"
		if int(circuit_average[x]["MAX THD"]) > 5:
			word = "does not fulfill"
		else:
			word = "fulfills"

	first = document.add_paragraph('As shown in the table, the rated current of the incomer ' + wording)
	first.add_run(' and the harmonic distortion percentage of current should not be above ' + circuitTHD + '. ')
	maxthd = str(circuit_average[x]["MAX THD"])
	first.add_run('The maximum current THD is at around ' + maxthd)
	first.add_run('%. Overall, the current harmonic distortion percentage of the measurement period ')
	first.add_run(word)
	first.add_run(' the requirement of the Code of Practice of Energy Efficiency of Building Services Installation.')

	document.add_paragraph('Figure presents the THD current result for ' + circuit_details[x]["Circuit Name"])

	filepath = os.path.abspath(
		'graphs/THD-F [%]_' + str(i) + '_' + circuit_details[x]["Circuit Name"] + '_' + circuit_details[x][
			"Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Current of THD-F of ' + circuit_details[x]["Circuit Name"] + ' for Measurement Period')

	document.add_paragraph()
	document.add_paragraph('Figure below shows a scatter plot between the phase currents and current total ' +
						   'harmonic distortion data points as measured during the logging period. The grey ' +
						   'shaded area represents the area that satisfies the maximum THD limit requirement ' +
						   'of the Code of Practice')
	filepath = os.path.abspath(
		'graphs/Current_vs_THD_' + str(i) + '_' + circuit_details[x]["Circuit Name"] + '_' + circuit_details[x][
			"Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Current and Phase Current Harmonic Distortion % of ' +
						   circuit_details[x]["Circuit Name"] + ' for Measurement Period')

	document.add_paragraph()
	p = document.add_paragraph().add_run('Power Factor')
	p.bold = True

	pfmin = str(circuit_average[x]["PF MIN"])
	pfmax = str(circuit_average[x]["PF MAX"])
	pfave = str(circuit_average[x]["AVE PF"])
	if circuit_average[x]["AVE PF"] < 0.85:
		words = "not as good as"
	else:
		words= "good as"
	document.add_paragraph('The average PF for ' + circuit_details[x]["Circuit Name"] + ' is ' +
						   pfave + '. The maximum power factor measurd is at ' +
						   pfmax + ' while the minimum is at ' + pfmin +
						   '. The ideal power factor is 1 and the acceptable power factor ranges from around 0.95 to' +
						   ' 0.85. This means that the performance of the cirsuit is ' + words +
						   ' the standard level.')

	# add voltage table and label format
	document.add_paragraph('. Power Factor of ' + circuit_details[x]["Circuit Name"] + ' for Measurement Period Data Summary')
	table1 = document.add_table(rows=2, cols=5)
	table1.style = 'Light List Accent 5'
	label_cell0 = table1.cell(0, 0)
	label0_paragraph = label_cell0.paragraphs[0]
	run0 = label0_paragraph.add_run('Parameter')
	run0.font.color.rgb = RGBColor(255, 255, 255)
	label_cell1 = table1.cell(0, 1)
	label1_paragraph = label_cell1.paragraphs[0]
	run1 = label1_paragraph.add_run('Times')
	run1.font.color.rgb = RGBColor(255, 255, 255)
	label_cell2 = table1.cell(0, 2)
	label2_paragraph = label_cell2.paragraphs[0]
	run2 = label2_paragraph.add_run('Power Factor Max.\n (Lagging Only)')
	run2.font.color.rgb = RGBColor(255, 255, 255)
	label_cell3 = table1.cell(0, 3)
	label3_paragraph = label_cell3.paragraphs[0]
	run3 = label3_paragraph.add_run('Power Factor Min.\n(Lagging Only)')
	run3.font.color.rgb = RGBColor(255, 255, 255)
	label_cell4 = table1.cell(0, 4)
	label4_paragraph = label_cell4.paragraphs[0]
	run4 = label4_paragraph.add_run('Power Factor Average (Lagging Only)')
	run4.font.color.rgb = RGBColor(255, 255, 255)

	# information of the circuit
	first_cells = table1.rows[1].cells
	first_cells[0].text = circuit_details[x]["Circuit Name"] + ' U1 [V]'
	first_cells[1].text = number_rows[x]
	curmax = str(circuit_details[x]["PF MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	first_cells[2].text = curmax
	curmin = str(circuit_details[x]["PF MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	first_cells[3].text = curmin
	curave = str(circuit_average[x]["AVE PF"])
	first_cells[4].text = curave
	document.add_paragraph()
	filepath = os.path.abspath(
		'graphs/Power factor_' + str(i) + '_' + circuit_details[x]["Circuit Name"] + '_' + circuit_details[x][
			"Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Power Factor of ' + circuit_details[x]["Circuit Name"] + ' for Measurement Period')

print("incomers done\n")
y = incomer_total - 1
for x in range(circuit_total - incomer_total):
	y = incomer_total + x
	document.add_heading(circuit_details[y]["Circuit Name"] + ' (' + circuit_details[y]["Circuit current"] + 'A)', level=1)
	p = document.add_paragraph().add_run('Current and Voltage')
	p.bold = True

	# add current table and label format
	document.add_paragraph('Current of ' + circuit_details[y]["Circuit Name"] + ' for Measurement Period Data Summary')
	table = document.add_table(rows=4, cols=5)
	table.style = 'Light List Accent 5'
	label_cell0 = table.cell(0, 0)
	label0_paragraph = label_cell0.paragraphs[0]
	run0 = label0_paragraph.add_run('Parameter')
	run0.font.color.rgb = RGBColor(255, 255, 255)
	label_cell1 = table.cell(0, 1)
	label1_paragraph = label_cell1.paragraphs[0]
	run1 = label1_paragraph.add_run('Times')
	run1.font.color.rgb = RGBColor(255, 255, 255)
	label_cell2 = table.cell(0, 2)
	label2_paragraph = label_cell2.paragraphs[0]
	run2 = label2_paragraph.add_run('Current RMS\n Max.')
	run2.font.color.rgb = RGBColor(255, 255, 255)
	label_cell3 = table.cell(0, 3)
	label3_paragraph = label_cell3.paragraphs[0]
	run3 = label3_paragraph.add_run('Current RMS\n Min.')
	run3.font.color.rgb = RGBColor(255, 255, 255)
	label_cell4 = table.cell(0, 4)
	label4_paragraph = label_cell4.paragraphs[0]
	run4 = label4_paragraph.add_run('Current RMS Average')
	run4.font.color.rgb = RGBColor(255, 255, 255)

	first_cells = table.rows[1].cells
	first_cells[0].text = circuit_details[y]["Circuit Name"] + ' I1 [A]'
	first_cells[1].text = number_rows[y]
	curmax = str(circuit_details[y]["CUR 1 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	first_cells[2].text = curmax
	curmin = str(circuit_details[y]["CUR 1 MIN"])
	first_cells[3].text = curmin
	curmin = curmin[1:]
	curmin = curmin[:-1]
	curave = str(circuit_average[y]["AVE i1"])
	first_cells[4].text = curave
	second_cells = table.rows[2].cells
	second_cells[0].text = circuit_details[y]["Circuit Name"] + ' I2 [A]'
	second_cells[1].text = number_rows[y]
	curmax = str(circuit_details[y]["CUR 2 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	second_cells[2].text = curmax
	curmin = str(circuit_details[y]["CUR 2 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	second_cells[3].text = curmin
	curave = str(circuit_average[y]["AVE i2"])
	second_cells[4].text = curave
	third_cells = table.rows[3].cells
	third_cells[0].text = circuit_details[y]["Circuit Name"] + ' I2 [A]'
	third_cells[1].text = number_rows[y]
	curmax = str(circuit_details[y]["CUR 3 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	third_cells[2].text = curmax
	curmin = str(circuit_details[y]["CUR 3 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	third_cells[3].text = curmin
	curave = str(circuit_average[y]["AVE i3"])
	third_cells[4].text = curave

	i = y+1
	filepath = os.path.abspath('graphs/Current [A]_' + str(i) + '_' + circuit_details[y]["Circuit Name"] + '_' +
							   circuit_details[y]["Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Current of ' + circuit_details[y]["Circuit Name"] + ' for Measurement Period ')
	document.add_paragraph()
	document.add_paragraph()

	document.add_paragraph('Voltage of ' + circuit_details[y]["Circuit Name"] + ' for Measurement Period Data Summary')
	table1 = document.add_table(rows=4, cols=5)
	table1.style = 'Light List Accent 5'
	label_cell0 = table1.cell(0, 0)
	label0_paragraph = label_cell0.paragraphs[0]
	run0 = label0_paragraph.add_run('Parameter')
	run0.font.color.rgb = RGBColor(255, 255, 255)
	label_cell1 = table1.cell(0, 1)
	label1_paragraph = label_cell1.paragraphs[0]
	run1 = label1_paragraph.add_run('Times')
	run1.font.color.rgb = RGBColor(255, 255, 255)
	label_cell2 = table1.cell(0, 2)
	label2_paragraph = label_cell2.paragraphs[0]
	run2 = label2_paragraph.add_run('Voltage RMS\n Max.')
	run2.font.color.rgb = RGBColor(255, 255, 255)
	label_cell3 = table1.cell(0, 3)
	label3_paragraph = label_cell3.paragraphs[0]
	run3 = label3_paragraph.add_run('Voltage RMS\n Min.')
	run3.font.color.rgb = RGBColor(255, 255, 255)
	label_cell4 = table1.cell(0, 4)
	label4_paragraph = label_cell4.paragraphs[0]
	run4 = label4_paragraph.add_run('Voltage RMS Average')
	run4.font.color.rgb = RGBColor(255, 255, 255)

	# information of the circuit
	first_cells = table1.rows[1].cells
	first_cells[0].text = circuit_details[y]["Circuit Name"] + ' U1 [V]'
	first_cells[1].text = number_rows[y]
	curmax = str(circuit_details[y]["VOL 1 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	first_cells[2].text = curmax
	curmin = str(circuit_details[y]["VOL 1 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	first_cells[3].text = curmin
	curave = str(circuit_average[y]["AVE v2"])
	first_cells[4].text = curave
	second_cells = table1.rows[2].cells
	second_cells[0].text = circuit_details[y]["Circuit Name"] + ' U2 [V]'
	second_cells[1].text = number_rows[y]
	curmax = str(circuit_details[y]["VOL 2 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	second_cells[2].text = curmax
	curmin = str(circuit_details[y]["VOL 2 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	second_cells[3].text = curmin
	curave = str(circuit_average[y]["AVE v2"])
	second_cells[4].text = curave
	third_cells = table1.rows[3].cells
	third_cells[0].text = circuit_details[y]["Circuit Name"] + ' U2 [V]'
	third_cells[1].text = number_rows[y]
	curmax = str(circuit_details[y]["VOL 3 MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	third_cells[2].text = curmax
	curmin = str(circuit_details[y]["VOL 3 MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	third_cells[3].text = curmin
	curave = str(circuit_average[y]["AVE v3"])
	third_cells[4].text = curave

	# add current of first voltage
	filepath = os.path.abspath(
		'graphs/Voltage [V]_' + str(i) + '_' + circuit_details[y]["Circuit Name"] + '_' + circuit_details[y][
			"Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Voltage of ' + circuit_details[x]["Circuit Name"] + ' for Measurement Period')

	document.add_paragraph()
	p = document.add_paragraph().add_run('Total Harmonic Distortion')
	p.bold = True

	filepath = os.path.abspath(
		'graphs/THD-F [%]_' + str(i) + '_' + circuit_details[y]["Circuit Name"] + '_' + circuit_details[y][
			"Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Current of THD-F of ' + circuit_details[x]["Circuit Name"] + ' for Measurement Period')

	document.add_paragraph()
	p = document.add_paragraph().add_run('Power Factor')
	p.bold = True

	document.add_paragraph('Power Factor of ' + circuit_details[x]["Circuit Name"] + ' for Measurement Period Data Summary')
	table1 = document.add_table(rows=2, cols=5)
	table1.style = 'Light List Accent 5'
	label_cell0 = table1.cell(0, 0)
	label0_paragraph = label_cell0.paragraphs[0]
	run0 = label0_paragraph.add_run('Parameter')
	run0.font.color.rgb = RGBColor(255, 255, 255)
	label_cell1 = table1.cell(0, 1)
	label1_paragraph = label_cell1.paragraphs[0]
	run1 = label1_paragraph.add_run('Times')
	run1.font.color.rgb = RGBColor(255, 255, 255)
	label_cell2 = table1.cell(0, 2)
	label2_paragraph = label_cell2.paragraphs[0]
	run2 = label2_paragraph.add_run('Power Factor Max.\n (Lagging Only)')
	run2.font.color.rgb = RGBColor(255, 255, 255)
	label_cell3 = table1.cell(0, 3)
	label3_paragraph = label_cell3.paragraphs[0]
	run3 = label3_paragraph.add_run('Power Factor Min.\n (Lagging Only)')
	run3.font.color.rgb = RGBColor(255, 255, 255)
	label_cell4 = table1.cell(0, 4)
	label4_paragraph = label_cell4.paragraphs[0]
	run4 = label4_paragraph.add_run('Power Factor Average (Lagging Only)')
	run4.font.color.rgb = RGBColor(255, 255, 255)

	# information of the circuit
	first_cells = table1.rows[1].cells
	first_cells[0].text = circuit_details[y]["Circuit Name"] + ' PF Sum Average'
	first_cells[1].text = number_rows[y]
	curmax = str(circuit_details[y]["PF MAX"])
	curmax = curmax[1:]
	curmax = curmax[:-1]
	first_cells[2].text = curmax
	curmin = str(circuit_details[y]["PF MIN"])
	curmin = curmin[1:]
	curmin = curmin[:-1]
	first_cells[3].text = curmin
	curave = str(circuit_average[y]["AVE PF"])
	first_cells[4].text = curave
	document.add_paragraph()
	filepath = os.path.abspath(
		'graphs/Power factor_' + str(i) + '_' + circuit_details[y]["Circuit Name"] + '_' + circuit_details[y][
			"Circuit current"] + '.png')
	document.add_picture(filepath, width=Inches(6.0))
	document.add_paragraph('Power Factor of ' + circuit_details[x]["Circuit Name"] +' for Measurement Periods')


print("All circuit is done\n")
document.add_heading('Summary of Power Quality Measurement Findings', level=1)
p = document.add_paragraph().add_run('Current and Voltage')
p.bold = True
document.add_paragraph('Table below summarizes the current RMS average and the voltage RMS average values as' +
					   ' measeured for the different phases in each circuit')

sum_row = circuit_total + 1
table1 = document.add_table(rows=sum_row, cols=7)
table1.style = 'Light List Accent 5'
label_cell0 = table1.cell(0, 0)
label0_paragraph = label_cell0.paragraphs[0]
run0 = label0_paragraph.add_run('Circuit')
run0.font.color.rgb = RGBColor(255, 255, 255)
label_cell1 = table1.cell(0, 1)
label1_paragraph = label_cell1.paragraphs[0]
run1 = label1_paragraph.add_run('I1 Current RMS Average (A)')
run1.font.color.rgb = RGBColor(255, 255, 255)
label_cell2 = table1.cell(0, 2)
label2_paragraph = label_cell2.paragraphs[0]
run2 = label2_paragraph.add_run('I2 Current RMS Average (A)')
run2.font.color.rgb = RGBColor(255, 255, 255)
label_cell3 = table1.cell(0, 3)
label3_paragraph = label_cell3.paragraphs[0]
run3 = label3_paragraph.add_run('I3 Current RMS Average (A)')
run3.font.color.rgb = RGBColor(255, 255, 255)
label_cell4 = table1.cell(0, 4)
label4_paragraph = label_cell4.paragraphs[0]
run4 = label4_paragraph.add_run('U1 Voltage RMS Average (V)')
run4.font.color.rgb = RGBColor(255, 255, 255)
label_cell5 = table1.cell(0, 5)
label5_paragraph = label_cell5.paragraphs[0]
run5 = label5_paragraph.add_run('U2 Voltage RMS Average (V)')
run5.font.color.rgb = RGBColor(255, 255, 255)
label_cell6 = table1.cell(0, 6)
label6_paragraph = label_cell6.paragraphs[0]
run6 = label6_paragraph.add_run('U3 Voltage RMS Average (V)')
run6.font.color.rgb = RGBColor(255, 255, 255)

for x in range(circuit_total):
	row_num = int(x + 1)
	first_cells = table1.rows[row_num].cells
	first_cells[0].text = circuit_details[x]["Circuit Name"]
	cur1 = str(circuit_average[x]["AVE i1"])
	first_cells[1].text = cur1
	cur2 = str(circuit_average[x]["AVE i2"])
	first_cells[2].text = cur2
	cur3 = str(circuit_average[x]["AVE i3"])
	first_cells[3].text = cur3
	vol1 = str(circuit_average[x]["AVE v1"])
	first_cells[4].text = vol1
	vol2 = str(circuit_average[x]["AVE v2"])
	first_cells[5].text = vol2
	vol3 = str(circuit_average[x]["AVE v3"])
	first_cells[6].text = vol3

document.add_paragraph()
p = document.add_paragraph().add_run('Total Harmonic Distortion')
p.bold = True
document.add_paragraph('Based on the measurements above, the maximum Total Harmonic Distortion (THD) in the ' +
					   'percentage of fundamental current for each circuit is summarized in Table below')

table1 = document.add_table(rows=sum_row, cols=7)
table1.style = 'Light List Accent 5'
label_cell0 = table1.cell(0, 0)
label0_paragraph = label_cell0.paragraphs[0]
run0 = label0_paragraph.add_run('Circuit')
run0.font.color.rgb = RGBColor(255, 255, 255)
label_cell1 = table1.cell(0, 1)
label1_paragraph = label_cell1.paragraphs[0]
run1 = label1_paragraph.add_run('THD-F % Current L1 Max.')
run1.font.color.rgb = RGBColor(255, 255, 255)
label_cell2 = table1.cell(0, 2)
label2_paragraph = label_cell2.paragraphs[0]
run2 = label2_paragraph.add_run('THD-F % Current L2 Max.')
run2.font.color.rgb = RGBColor(255, 255, 255)
label_cell3 = table1.cell(0, 3)
label3_paragraph = label_cell3.paragraphs[0]
run3 = label3_paragraph.add_run('THD-F % Current L3 Max.')
run3.font.color.rgb = RGBColor(255, 255, 255)
label_cell4 = table1.cell(0, 4)
label4_paragraph = label_cell4.paragraphs[0]
run4 = label4_paragraph.add_run('THD-F % Current L1 Avg.')
run4.font.color.rgb = RGBColor(255, 255, 255)
label_cell5 = table1.cell(0, 5)
label5_paragraph = label_cell5.paragraphs[0]
run5 = label5_paragraph.add_run('THD-F % Current L2 Avg.')
run5.font.color.rgb = RGBColor(255, 255, 255)
label_cell6 = table1.cell(0, 6)
label6_paragraph = label_cell6.paragraphs[0]
run6 = label6_paragraph.add_run('THD-F % Current L3 Avg.')
run6.font.color.rgb = RGBColor(255, 255, 255)

for x in range(circuit_total):
	row_num = int(x + 1)
	first_cells = table1.rows[row_num].cells
	first_cells[0].text = circuit_details[x]["Circuit Name"]
	cur1 = str(circuit_average[x]["MAX THD1"])
	first_cells[1].text = cur1
	cur2 = str(circuit_average[x]["MAX THD2"])
	first_cells[2].text = cur2
	cur3 = str(circuit_average[x]["MAX THD3"])
	first_cells[3].text = cur3
	vol1 = str(circuit_average[x]["AVE THD1"])
	first_cells[4].text = vol1
	vol2 = str(circuit_average[x]["AVE THD2"])
	first_cells[5].text = vol2
	vol3 = str(circuit_average[x]["AVE THD3"])
	first_cells[6].text = vol3

document.add_paragraph()
words1 = document.add_paragraph('As shown in the above table, the total harmonic distortion (THD) for ')
count = 0
count_and = 0
for x in range(circuit_total):
	if int(circuit_details[x]["Circuit current"])>= 400 and int(circuit_details[x]["Circuit current"]) < 800:
		if int(circuit_average[x]["MAX THD"]) > 12:
			count_and += 1
	elif int(circuit_details[x]["Circuit current"]) >= 800 and int(circuit_details[x]["Circuit current"]) < 2000:
		if int(circuit_average[x]["MAX THD"]) > 8:
			count_and += 1
	elif int(circuit_details[x]["Circuit current"]) > 2000:
		if int(circuit_average[x]["MAX THD"]) > 5:
			count_and += 1
for x in range(circuit_total):
	if int(circuit_details[x]["Circuit current"])>= 400 and int(circuit_details[x]["Circuit current"]) < 800:
		if int(circuit_average[x]["MAX THD"]) > 12:
			count += 1
			if count == count_and - 1:
				words1.add_run(circuit_details[x]["Circuit Name"] + ' and ')
			elif count == count_and:
				words1.add_run(circuit_details[x]["Circuit Name"] +' ')
			else:
				words1.add_run(circuit_details[x]["Circuit Name"] + ', ')
	elif int(circuit_details[x]["Circuit current"]) >= 800 and int(circuit_details[x]["Circuit current"]) < 2000:
		if int(circuit_average[x]["MAX THD"]) > 8:
			count += 1
			if count == count_and - 1:
				words1.add_run(circuit_details[x]["Circuit Name"] + ' and ')
			elif count == count_and:
				words1.add_run(circuit_details[x]["Circuit Name"] +' ')
			else:
				words1.add_run(circuit_details[x]["Circuit Name"] + ', ')
	elif int(circuit_details[x]["Circuit current"]) > 2000:
		if int(circuit_average[x]["MAX THD"]) > 5:
			count += 1
			if count == count_and - 1:
				words1.add_run(circuit_details[x]["Circuit Name"] + ' and ')
			elif count == count_and:
				words1.add_run(circuit_details[x]["Circuit Name"] + ' ')
			else:
				words1.add_run(circuit_details[x]["Circuit Name"] + ', ')

if count == 0:
	words1.add_run('all circuits satisfy the minimum THD requirement of the Building Energy Code.')
elif count == 1:
	words1.add_run('is greater that the minimum THD requirement of the Building Energy Code.')
elif count > 1:
	words1.add_run('are greater that the minimum THD requirement of the Building Energy Code.')
words1.add_run(' For detailed power quality log data, please refer to Appendix D,')
document.add_paragraph()
filepath = os.path.abspath('graphs/THD' + '.png')
document.add_picture(filepath, width=Inches(4.0))
document.add_paragraph('Code of Practice of Energy Efficiency of Building Services Installation')

document.add_paragraph()
p = document.add_paragraph().add_run('Power Factor')
p.bold = True
words2 = document.add_paragraph('Table below summarizes the measured power factor for each circuit. Overall, ')
count = 0
for x in range(circuit_total):
	if circuit_average[x]["AVE PF"] < 0.85:
		words2.add_run(circuit_details[x]["Circuit Name"] + ', ')
		count += 1

if count == 0:
	words2.add_run('all circuits have standard power factors.')
elif count == 1:
	words2.add_run('falls below the minimum recommended power factor of 0.85.')
elif count > 1:
	words2.add_run('fall below the minimum recommended power factor of 0.85.')

document.add_paragraph()

table1 = document.add_table(rows=sum_row, cols=4)
table1.style = 'Light List Accent 5'
label_cell0 = table1.cell(0, 0)
label0_paragraph = label_cell0.paragraphs[0]
run0 = label0_paragraph.add_run('Circuit')
run0.font.color.rgb = RGBColor(255, 255, 255)
label_cell1 = table1.cell(0, 1)
label1_paragraph = label_cell1.paragraphs[0]
run1 = label1_paragraph.add_run('Average Power Factor')
run1.font.color.rgb = RGBColor(255, 255, 255)
label_cell2 = table1.cell(0, 2)
label2_paragraph = label_cell2.paragraphs[0]
run2 = label2_paragraph.add_run('Maximum Power Factor')
run2.font.color.rgb = RGBColor(255, 255, 255)
label_cell3 = table1.cell(0, 3)
label3_paragraph = label_cell3.paragraphs[0]
run3 = label3_paragraph.add_run('Minimum Power Factor')
run3.font.color.rgb = RGBColor(255, 255, 255)

for x in range(circuit_total):
	row_num = int(x + 1)
	first_cells = table1.rows[row_num].cells
	first_cells[0].text = circuit_details[x]["Circuit Name"]
	cur1 = str(circuit_average[x]["AVE PF"])
	first_cells[1].text = cur1
	cur2 = str(circuit_average[x]["PF MAX"])
	first_cells[2].text = cur2
	cur3 = str(circuit_average[x]["PF MIN"])
	first_cells[3].text = cur3

document.save('Section 5 Power Quality Analysis.docx')

power_quality_appendix(circuit_details, circuit_total)
